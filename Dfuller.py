# I use this script to do a Dickey-Fuller test to determine if my variablees are stable or need be differenced.


import os
import numpy as np
import pandas as pd
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from docx import Document

def select_best_lag(series: pd.Series, p_max: int) -> int:
    dy = series.diff().dropna()

    best_bic = np.inf
    best_aic = np.inf
    best_lag = 0

    for lag in range(p_max, -1, -1):
        dfX = pd.DataFrame({
            'y_lag1': series.shift(1)
        }).loc[dy.index]

        for i in range(1, lag+1):
            dfX[f'dy_lag_{i}'] = dy.shift(i)

        df_full = pd.concat([dy, dfX], axis=1).dropna()
        y_reg = df_full.iloc[:, 0]
        X_reg = sm.add_constant(df_full.iloc[:, 1:])

        res = sm.OLS(y_reg, X_reg).fit()

        if res.bic < best_bic - 1e-8:
            best_bic, best_aic, best_lag = res.bic, res.aic, lag
        elif abs(res.bic - best_bic) < 1e-8 and res.aic < best_aic:
            best_aic, best_lag = res.aic, lag

    return best_lag

def main():
    desktop    = os.path.expanduser("~/Desktop")
    excel_fp   = os.path.join(desktop, "CMC - Dfuller.xlsx")
    report_fp  = os.path.join(desktop, "ADF_Test_Results.docx")

    df = pd.read_excel(excel_fp, sheet_name="Sheet1")

    variables = [
        "Order imbalances",
        "Annualized SD",
        "Liq - price impact",
        "Liq - dept"
    ]

    df.dropna(subset=variables, inplace=True)

    summary = {}

    for var in variables:
        if var not in df.columns:
            print(f"⚠️  '{var}' not found in data; skipping.")
            continue

        series = df[var]
        n = len(series)
        p_max = int(12 * ((n / 100) ** 0.25))
        p_max = max(p_max, 0)

        best_lag = select_best_lag(series, p_max)

        adf_res = adfuller(series, maxlag=best_lag, regression='c', autolag=None)
        adf_stat, p_value, used_lag, nobs, crit_vals = adf_res[:5]

        summary[var] = {
            "n_obs":    int(nobs),
            "lag_used": int(used_lag),
            "adf_stat": float(adf_stat),
            "p_value":  float(p_value),
            "crit_vals": {k: float(v) for k, v in crit_vals.items()}
        }

        print(f"{var}: n={n}, p_max={p_max}, selected lag={best_lag}, "
              f"ADF={adf_stat:.3f}, p={p_value:.4f}")

    doc = Document()
    doc.add_heading("Augmented Dickey-Fuller Test Results", level=0)

    for var, res in summary.items():
        doc.add_heading(var, level=1)
        doc.add_paragraph(f"Number of obs = {res['n_obs']},   "
                          f"Number of lags = {res['lag_used']}")
        doc.add_paragraph("H0: The series has a unit root (random walk without drift).")
        doc.add_paragraph("Test conducted at 5% significance level.")
        doc.add_paragraph(f"ADF Test statistic = {res['adf_stat']:.3f}")
        cv = res["crit_vals"]
        doc.add_paragraph(
            f"Critical values: 1% = {cv['1%']:.3f}, "
            f"5% = {cv['5%']:.3f}, 10% = {cv['10%']:.3f}"
        )
        doc.add_paragraph(f"MacKinnon approximate p-value = {res['p_value']:.4f}")

    doc.save(report_fp)
    print(f"\n✅ Word report written to:\n   {report_fp}")

if __name__ == "__main__":
    main()
